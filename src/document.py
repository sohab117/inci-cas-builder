"""Generate the regulatory CAS reference .docx from looked-up panel entries.

Phase 1.3.5: landscape orientation, fixed column widths, simplified function
field for the Function column (full CosIng list preserved as `function_full`),
darker yellow on Verified=N rows, and rows that won't split across pages.
The Trade Name / Source column was dropped — supplier data isn't yet wired in
and the empty column was just absorbing horizontal space.
"""

from __future__ import annotations

from datetime import date as _date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

# Palette
HEADER_FILL = "D5E8F0"  # blue-tinted header
UNVERIFIED_FILL = "FFE699"  # warmer yellow for "needs review" rows
NOTE_BORDER_COLOR = "F0C419"  # yellow accent on the "Note:" paragraph
META_GRAY = RGBColor(0x55, 0x55, 0x55)

DASH = "—"

COLUMN_HEADERS = [
    "#",
    "INCI Name",
    "Common Name / Function",
    "CAS Number",
    "EINECS",
    "Verified",
]

# Inches — landscape Letter with 0.5" margins gives ~10" usable; total ~8.8"
COLUMN_WIDTHS_INCHES = [0.4, 2.2, 2.6, 1.4, 1.4, 0.8]

VERIFIED_COL_INDEX = len(COLUMN_HEADERS) - 1  # last column

DEFAULT_NOTE = (
    "Note: Rows marked N in the Verified column require confirmation against "
    "supplier SDS or alternative sources before regulatory submission."
)


def simplify_function(function_str: str | None, inci_name: str = "") -> str:
    """Pick a single concise category from CosIng's verbose Function list.

    Priority order:
      1. Surfactant types (sub-categorized via INCI name keywords)
      2. Preservative
      3. Solvent
      4. Skin Conditioning - Emollient -> "Emollient"
      5. Humectant
      6. Antimicrobial -> "Preservative"
      7. Viscosity Controlling -> "Viscosity Modifier"
      8. Plain "Skin Conditioning"
      9. Fall back to the first item in the CosIng list
    """
    if not function_str:
        return ""

    fn = function_str.lower()
    name = (inci_name or "").lower()

    if "surfactant" in fn:
        if any(kw in name for kw in ("sodium", "sulfonate", "isethionate")):
            return "Anionic Surfactant"
        if "betaine" in name:
            return "Amphoteric Surfactant"
        if any(kw in name for kw in ("glucoside", "glyceride")):
            return "Nonionic Surfactant"
        return "Surfactant"

    if "preservative" in fn:
        return "Preservative"

    if "solvent" in fn:
        return "Solvent"

    if "emollient" in fn:
        return "Emollient"

    if "humectant" in fn:
        return "Humectant"

    if "antimicrobial" in fn:
        return "Preservative"

    if "viscosity" in fn:
        return "Viscosity Modifier"

    if "skin conditioning" in fn:
        return "Skin Conditioning"

    return function_str.split(",")[0].strip()


def generate_document(
    parsed_entries: list[dict],
    output_path: str,
    metadata: dict | None = None,
) -> str:
    """Render parsed+looked-up entries to a .docx file. Returns the path."""
    metadata = metadata or {}
    product_name = metadata.get("product_name") or "Cosmetic Product"
    client_name = metadata.get("client_name") or ""
    prepared_by = metadata.get("prepared_by") or ""
    purpose = metadata.get("purpose") or ""
    doc_date = metadata.get("date") or _date.today().strftime("%Y-%m-%d")

    doc = Document()
    _configure_landscape_letter(doc)

    _write_header(doc, product_name, prepared_by, client_name, doc_date, purpose)

    footnotes = _FootnoteLedger()
    _write_table(doc, parsed_entries, footnotes)

    _write_footnotes_section(doc, footnotes)
    _write_confidentiality(doc, prepared_by)

    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    return output_path


# ---------------------------------------------------------------------------
# Document sections
# ---------------------------------------------------------------------------


def _configure_landscape_letter(doc) -> None:
    """Set landscape US Letter with 0.5" margins."""
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    # python-docx does not auto-swap dimensions when toggling orientation;
    # we set them explicitly so the page is actually wider than tall.
    section.page_width = Inches(11.0)
    section.page_height = Inches(8.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)


def _write_header(doc, product_name, prepared_by, client_name, doc_date, purpose):
    title_para = doc.add_paragraph()
    title_run = title_para.add_run(f"{product_name} {DASH} CAS Number Reference")
    title_run.bold = True
    title_run.font.size = Pt(16)
    title_run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    title_para.paragraph_format.space_after = Pt(6)

    meta_parts: list[str] = []
    if prepared_by:
        meta_parts.append(f"Prepared by: {prepared_by}")
    if client_name:
        meta_parts.append(f"Client: {client_name}")
    meta_parts.append(f"Date: {doc_date}")
    meta_para = doc.add_paragraph()
    meta_run = meta_para.add_run("  |  ".join(meta_parts))
    meta_run.font.size = Pt(10)
    meta_run.font.color.rgb = META_GRAY

    if purpose:
        purpose_para = doc.add_paragraph()
        purpose_run = purpose_para.add_run(f"Purpose: {purpose}")
        purpose_run.font.size = Pt(10)
        purpose_run.italic = True
        purpose_run.font.color.rgb = META_GRAY

    note_para = doc.add_paragraph()
    note_para.paragraph_format.left_indent = Pt(12)
    bold_label = note_para.add_run("Note: ")
    bold_label.bold = True
    bold_label.font.size = Pt(10)
    body_run = note_para.add_run(DEFAULT_NOTE.removeprefix("Note: "))
    body_run.font.size = Pt(10)
    _add_left_border(note_para, color=NOTE_BORDER_COLOR, size_eighths_pt=24, space=6)


def _write_table(doc, entries, footnotes):
    table = doc.add_table(rows=1, cols=len(COLUMN_HEADERS))
    table.style = "Table Grid"
    table.autofit = False
    table.allow_autofit = False
    _set_table_layout_fixed(table)

    header_row = table.rows[0]
    header_row.cant_split = True
    for i, header in enumerate(COLUMN_HEADERS):
        cell = header_row.cells[i]
        cell.width = Inches(COLUMN_WIDTHS_INCHES[i])
        cell.text = ""
        run = cell.paragraphs[0].add_run(header)
        run.bold = True
        _shade_cell(cell, HEADER_FILL)

    for entry in entries:
        row = table.add_row()
        row.cant_split = True
        _write_data_row(row, entry, footnotes)


def _write_data_row(row, entry, footnotes):
    cells = row.cells
    for i, width in enumerate(COLUMN_WIDTHS_INCHES):
        cells[i].width = Inches(width)

    cells[0].text = str(entry.get("position", ""))

    inci_cell = cells[1]
    inci_cell.text = ""
    inci_cell.paragraphs[0].add_run(entry.get("inci_name") or "")
    synonyms = entry.get("synonyms") or []
    if synonyms:
        synonym_para = inci_cell.add_paragraph()
        joined = " / ".join([entry.get("inci_name") or ""] + synonyms)
        run = synonym_para.add_run(joined)
        run.italic = True

    cells[2].text = simplify_function(
        entry.get("function") or "",
        entry.get("inci_name") or "",
    )
    cells[3].text = entry.get("cas_number") or DASH
    cells[4].text = entry.get("einecs_number") or DASH

    verified_cell = cells[VERIFIED_COL_INDEX]
    verified_cell.text = ""
    verified_para = verified_cell.paragraphs[0]
    verified = bool(entry.get("verified"))
    label = verified_para.add_run("Y" if verified else "N")
    label.bold = True

    note = entry.get("verification_note")
    if note:
        ref = verified_para.add_run(str(footnotes.ref(note)))
        ref.font.superscript = True
        ref.bold = True

    if not verified:
        for cell in cells:
            _shade_cell(cell, UNVERIFIED_FILL)


def _write_footnotes_section(doc, footnotes):
    if not footnotes:
        return
    doc.add_paragraph()
    for num, text in footnotes.ordered():
        para = doc.add_paragraph()
        ref = para.add_run(str(num))
        ref.font.superscript = True
        body = para.add_run(" " + text)
        body.italic = True
        for run in (ref, body):
            run.font.size = Pt(10)


def _write_confidentiality(doc, prepared_by):
    para = doc.add_paragraph()
    suffix = f" | {prepared_by}" if prepared_by else ""
    run = para.add_run(
        f"Confidential {DASH} For Internal and Regulatory Use Only{suffix}"
    )
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = META_GRAY


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


class _FootnoteLedger:
    """Collects unique verification_note strings and assigns 1-indexed numbers."""

    def __init__(self) -> None:
        self._notes: dict[str, int] = {}

    def __bool__(self) -> bool:
        return bool(self._notes)

    def ref(self, note: str) -> int:
        if note not in self._notes:
            self._notes[note] = len(self._notes) + 1
        return self._notes[note]

    def ordered(self) -> list[tuple[int, str]]:
        return sorted(((n, t) for t, n in self._notes.items()), key=lambda x: x[0])


def _shade_cell(cell, hex_color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def _set_table_layout_fixed(table) -> None:
    """Word's "fixed" table layout — columns honor explicit widths."""
    tbl_pr = table._tbl.tblPr
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tbl_pr.append(layout)


def _add_left_border(paragraph, *, color: str, size_eighths_pt: int, space: int) -> None:
    """Render a colored left border on a paragraph (the 'Note:' rule)."""
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), str(size_eighths_pt))
    left.set(qn("w:space"), str(space))
    left.set(qn("w:color"), color)
    p_bdr.append(left)
    p_pr.append(p_bdr)
