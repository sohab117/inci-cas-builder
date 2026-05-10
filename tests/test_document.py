"""Unit tests for src/document.py — .docx generation."""

import pytest
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn

from src.document import generate_document, simplify_function


def _entry(
    position,
    inci_name,
    synonyms=None,
    function=None,
    cas_number=None,
    einecs_number=None,
    verified=False,
    source="not_found",
    verification_note=None,
):
    """Build a lookup-style entry dict for tests."""
    return {
        "position": position,
        "inci_name": inci_name,
        "inci_normalized": inci_name.upper(),
        "raw": inci_name,
        "synonyms": synonyms or [],
        "notes": [],
        "function": function,
        "cas_number": cas_number,
        "einecs_number": einecs_number,
        "verified": verified,
        "source": source,
        "verification_note": verification_note,
    }


def _cell_fill(cell):
    """Return the hex fill color of a table cell, or None."""
    tcPr = cell._tc.tcPr
    if tcPr is None:
        return None
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        return None
    return shd.get(qn("w:fill"))


def _find_data_table(doc):
    """The document has exactly one table (the ingredient table). Return it."""
    assert len(doc.tables) == 1, f"expected 1 table, found {len(doc.tables)}"
    return doc.tables[0]


def test_generates_valid_docx_that_reopens(tmp_path):
    out = tmp_path / "out.docx"
    entries = [_entry(1, "Water", cas_number="7732-18-5", verified=True, source="cosing")]

    returned = generate_document(entries, str(out))

    assert returned == str(out)
    assert out.exists() and out.stat().st_size > 0
    Document(str(out))


def test_table_has_6_columns_and_correct_row_count(tmp_path):
    out = tmp_path / "out.docx"
    entries = [
        _entry(1, "Water", cas_number="7732-18-5", verified=True, source="cosing"),
        _entry(2, "Glycerin", cas_number="56-81-5", verified=True, source="cosing"),
        _entry(3, "Tocopherol", verified=False, source="not_found"),
    ]

    generate_document(entries, str(out))
    doc = Document(str(out))
    table = _find_data_table(doc)

    assert len(table.columns) == 6
    assert len(table.rows) == 4  # 1 header + 3 data

    header_texts = [c.text.strip() for c in table.rows[0].cells]
    expected = [
        "#",
        "INCI Name",
        "Common Name / Function",
        "CAS Number",
        "EINECS",
        "Verified",
    ]
    assert header_texts == expected


def test_table_has_no_trade_name_source_column(tmp_path):
    out = tmp_path / "out.docx"
    entries = [_entry(1, "Water", cas_number="7732-18-5", verified=True, source="cosing")]
    generate_document(entries, str(out))
    doc = Document(str(out))
    table = _find_data_table(doc)
    headers = [c.text.strip().lower() for c in table.rows[0].cells]
    assert not any("trade name" in h for h in headers)


def test_document_is_landscape_letter(tmp_path):
    out = tmp_path / "out.docx"
    entries = [_entry(1, "Water", cas_number="7732-18-5", verified=True, source="cosing")]
    generate_document(entries, str(out))
    doc = Document(str(out))
    section = doc.sections[0]
    assert section.orientation == WD_ORIENT.LANDSCAPE
    assert section.page_width > section.page_height, (
        f"page_width ({section.page_width}) must exceed page_height ({section.page_height})"
    )


def test_header_metadata_appears_in_document(tmp_path):
    out = tmp_path / "out.docx"
    entries = [_entry(1, "Water", cas_number="7732-18-5", verified=True, source="cosing")]
    metadata = {
        "product_name": "Vitasana Body Wash",
        "client_name": "Vitasana",
        "prepared_by": "117 Holdings LLC",
        "purpose": "EWG Verified submission support",
        "date": "2026-05-09",
    }

    generate_document(entries, str(out), metadata)
    doc = Document(str(out))

    body_text = "\n".join(p.text for p in doc.paragraphs)
    assert "Vitasana Body Wash" in body_text
    assert "Vitasana" in body_text
    assert "117 Holdings LLC" in body_text
    assert "EWG Verified submission support" in body_text
    assert "2026-05-09" in body_text


def test_unverified_rows_have_different_shading_than_verified_rows(tmp_path):
    out = tmp_path / "out.docx"
    entries = [
        _entry(1, "Water", cas_number="7732-18-5", verified=True, source="cosing"),
        _entry(2, "Mystery", verified=False, source="not_found"),
    ]

    generate_document(entries, str(out))
    doc = Document(str(out))
    table = _find_data_table(doc)

    verified_row_fill = _cell_fill(table.rows[1].cells[0])
    unverified_row_fill = _cell_fill(table.rows[2].cells[0])
    assert verified_row_fill != unverified_row_fill


def test_verification_notes_render_as_footnotes(tmp_path):
    out = tmp_path / "out.docx"
    entries = [
        _entry(1, "Water", cas_number="7732-18-5", verified=True, source="cosing"),
        _entry(
            2,
            "Sodium Lauroyl Methyl Isethionate",
            function="Cleansing",
            verified=False,
            source="cosing_partial",
            verification_note="CosIng entry exists but CAS field is empty in source data",
        ),
    ]

    generate_document(entries, str(out))
    doc = Document(str(out))
    table = _find_data_table(doc)

    # Verified column is now the last column (index 5 with 6 columns)
    verified_cell = table.rows[2].cells[-1]
    superscript_runs = [
        r for p in verified_cell.paragraphs for r in p.runs if r.font.superscript
    ]
    assert len(superscript_runs) >= 1, "no superscript footnote ref in unverified row"

    body_text = "\n".join(p.text for p in doc.paragraphs)
    assert "CosIng entry exists but CAS field is empty in source data" in body_text


def test_synonyms_render_as_italic_second_line_under_inci_name(tmp_path):
    out = tmp_path / "out.docx"
    entries = [
        _entry(
            1,
            "Aqua",
            synonyms=["Water", "Eau"],
            cas_number="7732-18-5",
            verified=True,
            source="cosing",
        ),
    ]

    generate_document(entries, str(out))
    doc = Document(str(out))
    table = _find_data_table(doc)

    inci_cell = table.rows[1].cells[1]
    paragraphs = inci_cell.paragraphs
    assert len(paragraphs) >= 2
    assert "Aqua" in paragraphs[0].text
    second_text = paragraphs[1].text
    assert "Water" in second_text and "Eau" in second_text
    assert any(r.italic for r in paragraphs[1].runs)


# ---------------------------------------------------------------------------
# simplify_function rules
# ---------------------------------------------------------------------------


@pytest.mark.parametrize(
    "function_str, inci_name, expected",
    [
        # Surfactant sub-categorization based on INCI name keywords
        ("Surfactant - Cleansing", "Cocamidopropyl Betaine", "Amphoteric Surfactant"),
        (
            "Surfactant - Cleansing",
            "Sodium Lauroyl Methyl Isethionate",
            "Anionic Surfactant",
        ),
        ("Surfactant - Cleansing", "Lauryl Glucoside", "Nonionic Surfactant"),
        ("Cleansing, Surfactant - Foam Boosting", "Disodium Lauroamphodiacetate", "Anionic Surfactant"),
        ("Surfactant - Emulsifying", "Mystery Compound", "Surfactant"),
        # Preservative & related
        ("Antimicrobial, Preservative", "", "Preservative"),
        ("Preservative", "", "Preservative"),
        ("Antimicrobial", "", "Preservative"),
        # Solvent
        ("Solvent", "", "Solvent"),
        ("Antiplaque, Skin Conditioning, Solvent", "Water", "Solvent"),
        # Skin Conditioning variants
        (
            "Deodorant, Hair Conditioning, Skin Conditioning, Skin Conditioning - Emollient",
            "Caprylyl Glycol",
            "Emollient",
        ),
        ("Skin Conditioning", "Aloe Barbadensis Leaf Juice", "Skin Conditioning"),
        # Other categories
        ("Humectant", "", "Humectant"),
        ("Viscosity Controlling", "", "Viscosity Modifier"),
        # Edge cases
        (None, "", ""),
        ("", "", ""),
        # Fall-through to first item
        ("Antistatic", "", "Antistatic"),
        ("Antistatic, Hair Conditioning", "", "Antistatic"),
    ],
)
def test_simplify_function_rules(function_str, inci_name, expected):
    assert simplify_function(function_str, inci_name) == expected
