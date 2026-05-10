"""Tests for the FastAPI surface in api/."""

import base64
import io

import pytest
from docx import Document
from fastapi.testclient import TestClient

import src.lookup as lookup_mod
from api.analyze import app as analyze_app
from api.health import app as health_app

VITASANA_PANEL = (
    "Water, Sodium Lauroyl Methyl Isethionate, Cocamidopropyl Betaine, "
    "Lauryl Glucoside, Aloe Barbadensis Leaf Juice*, Phenoxyethanol, "
    "Caprylyl Glycol, Ethylhexylglycerin"
)


@pytest.fixture(autouse=True)
def isolated_lookup(tmp_path, monkeypatch):
    monkeypatch.setattr(lookup_mod, "CACHE_PATH", tmp_path / "cache.db")
    monkeypatch.setattr(lookup_mod, "_cosing_cache", None)
    monkeypatch.delenv("ANTHROPIC_API_KEY", raising=False)
    yield


@pytest.fixture
def analyze_client():
    return TestClient(analyze_app)


@pytest.fixture
def health_client():
    return TestClient(health_app)


def test_health_returns_ok(health_client):
    resp = health_client.get("/api/health")
    assert resp.status_code == 200
    body = resp.json()
    assert body["status"] == "ok"
    assert "version" in body


def test_analyze_vitasana_panel_returns_expected_summary(analyze_client, mocker):
    pubchem_404 = mocker.MagicMock(status_code=404)
    mocker.patch("src.lookup.requests.get", return_value=pubchem_404)

    resp = analyze_client.post(
        "/api/analyze",
        json={
            "inci_string": VITASANA_PANEL,
            "metadata": {"product_name": "Vitasana Body Wash"},
        },
    )
    assert resp.status_code == 200, resp.text
    body = resp.json()

    assert len(body["ingredients"]) == 8
    assert body["summary"] == {
        "total": 8,
        "fully_verified": 7,
        "partial": 1,
        "not_found": 0,
    }
    assert isinstance(body["document_base64"], str)
    assert len(body["document_base64"]) > 100

    # Spot-check ingredient shape — `function` should be simplified, `function_full`
    # should preserve the full CosIng list.
    water = body["ingredients"][0]
    assert water["inci_name"] == "Water"
    assert water["cas_number"] == "7732-18-5"
    assert water["verified"] is True
    assert water["source"] == "cosing"
    assert water["function"] == "Solvent"
    assert water["function_full"] is not None and "Solvent" in water["function_full"]
    assert water["function_full"] != water["function"], (
        "function_full should still hold the full CosIng list, not the simplified one"
    )

    # Cocamidopropyl Betaine -> Amphoteric Surfactant via name keyword
    cocamido = next(
        i for i in body["ingredients"] if i["inci_name"] == "Cocamidopropyl Betaine"
    )
    assert cocamido["function"] == "Amphoteric Surfactant"


def test_analyze_returns_decodable_docx(analyze_client, mocker):
    pubchem_404 = mocker.MagicMock(status_code=404)
    mocker.patch("src.lookup.requests.get", return_value=pubchem_404)

    resp = analyze_client.post(
        "/api/analyze",
        json={"inci_string": "Water, Glycerin"},
    )
    assert resp.status_code == 200

    doc_bytes = base64.b64decode(resp.json()["document_base64"])
    doc = Document(io.BytesIO(doc_bytes))
    assert len(doc.tables) == 1
    table = doc.tables[0]
    # 1 header + 2 data rows
    assert len(table.rows) == 3
    assert len(table.columns) == 6


def test_analyze_empty_inci_string_returns_400(analyze_client):
    resp = analyze_client.post("/api/analyze", json={"inci_string": ""})
    assert resp.status_code == 400
    assert "empty" in resp.json()["detail"].lower()


def test_analyze_malformed_json_returns_422(analyze_client):
    resp = analyze_client.post("/api/analyze", json={"wrong_field": "blah"})
    assert resp.status_code == 422
