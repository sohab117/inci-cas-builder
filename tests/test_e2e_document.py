"""Full pipeline integration test: parse_inci -> lookup_panel -> generate_document.

PubChem is mocked to 404 (cosing_partial fallthrough exercised in test_e2e_lookup;
this test only verifies the rendered .docx structure). LLM disabled via fixture.

Output: output/test_vitasana.docx — gitignored, kept on disk for manual inspection.
"""

from pathlib import Path

import pytest
from docx import Document

import src.lookup as lookup_mod
from src.document import generate_document
from src.lookup import lookup_panel
from src.parser import parse_inci

VITASANA_PANEL = (
    "Water, Sodium Lauroyl Methyl Isethionate, Cocamidopropyl Betaine, "
    "Lauryl Glucoside, Aloe Barbadensis Leaf Juice*, Phenoxyethanol, "
    "Caprylyl Glycol, Ethylhexylglycerin"
)

OUTPUT_PATH = Path(__file__).resolve().parent.parent / "output" / "test_vitasana.docx"


@pytest.fixture(autouse=True)
def isolated_lookup(tmp_path, monkeypatch):
    monkeypatch.setattr(lookup_mod, "CACHE_PATH", tmp_path / "cache.db")
    monkeypatch.setattr(lookup_mod, "_cosing_cache", None)
    monkeypatch.delenv("ANTHROPIC_API_KEY", raising=False)
    yield


def test_full_pipeline_renders_vitasana_panel(mocker, capsys):
    pubchem_404 = mocker.MagicMock(status_code=404)
    mocker.patch("src.lookup.requests.get", return_value=pubchem_404)

    parsed = parse_inci(VITASANA_PANEL)
    looked_up = lookup_panel(parsed)
    metadata = {
        "product_name": "Vitasana Body Wash",
        "client_name": "Vitasana",
        "prepared_by": "117 Holdings LLC",
        "purpose": "EWG Verified submission support",
    }

    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    path = generate_document(looked_up, str(OUTPUT_PATH), metadata)

    print(f"\nGenerated: {path}")

    doc = Document(path)
    assert len(doc.tables) == 1
    table = doc.tables[0]
    # 1 header row + 8 ingredient rows
    assert len(table.rows) == 9, f"expected 9 rows (1 header + 8 data), got {len(table.rows)}"
    assert len(table.columns) == 6
